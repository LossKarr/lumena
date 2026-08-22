from __future__ import annotations

import json
from pathlib import Path

from docx import Document
import pytest

from src.documents.studio import DocumentStudio
from src.reasoning.handlers.documents import (
    apply_document_edit_handler,
    convert_library_document_handler,
    export_library_document_handler,
    get_document_history_handler,
    preview_document_edit_handler,
    revise_studio_document_handler,
)


ROOT = Path(__file__).parents[2]


@pytest.fixture
def studio(tmp_path, monkeypatch):
    instance = DocumentStudio(
        tmp_path / "studio",
        builtin_root=ROOT / "assets" / "templates",
        output_root=tmp_path / "output",
    )
    monkeypatch.setattr("src.documents.studio.get_document_studio", lambda: instance)
    return instance


@pytest.mark.asyncio
async def test_agent_can_revise_track_convert_and_export_studio_document(studio):
    template = studio.catalog.get("devis")
    generated = await studio.generate(
        kind="devis",
        output_format="html",
        data=studio.catalog.read_sample_data(template),
        filename="devis-agent",
    )
    document_id = generated["record"]["id"]

    revised = await revise_studio_document_handler(
        None, document_id, json.dumps({"numero": "AGENT-2026"})
    )
    assert revised.success
    child = json.loads(revised.output)["record"]
    assert child["parent_id"] == document_id

    history = await get_document_history_handler(None, document_id)
    assert history.success
    assert json.loads(history.output)["transformations"][0]["operation"] == "revise"

    converted = await convert_library_document_handler(None, child["id"], "pdf")
    assert converted.success
    converted_record = json.loads(converted.output)["record"]
    assert converted_record["format"] == "pdf"

    exported = await export_library_document_handler(None, converted_record["id"], "preuve-agent.pdf")
    assert exported.success
    assert Path(json.loads(exported.output)["proof"]).is_file()


@pytest.mark.asyncio
async def test_agent_lifecycle_accepts_an_exact_filename_reference(studio):
    template = studio.catalog.get("devis")
    generated = await studio.generate(
        kind="devis",
        output_format="html",
        data=studio.catalog.read_sample_data(template),
        filename="devis-par-nom",
    )
    original = generated["record"]

    revised = await revise_studio_document_handler(
        None, original["filename"], json.dumps({"numero": "NOM-2026"})
    )
    assert revised.success
    child = json.loads(revised.output)["record"]
    assert child["parent_id"] == original["id"]

    history = await get_document_history_handler(None, original["filename"])
    assert history.success
    assert json.loads(history.output)["document"]["id"] == original["id"]


@pytest.mark.asyncio
async def test_agent_office_edit_is_previewed_then_applied_transactionally(studio, tmp_path):
    source = tmp_path / "contrat.docx"
    document = Document()
    document.add_paragraph("Montant ancien")
    document.save(source)
    record, _ = studio.importer.import_file(source, source_kind="manual")
    operations = json.dumps([{"op": "replace_text", "find": "ancien", "replace": "nouveau"}])

    preview = await preview_document_edit_handler(None, record.id, operations)
    assert preview.success
    assert len(json.loads(preview.output)["operations"]) == 1

    applied = await apply_document_edit_handler(None, record.id, operations)
    assert applied.success
    child = json.loads(applied.output)["record"]
    assert child["parent_id"] == record.id
    assert "nouveau" in " ".join(p.text for p in Document(child["path"]).paragraphs)
    assert "ancien" in " ".join(p.text for p in Document(source).paragraphs)


@pytest.mark.asyncio
async def test_agent_rejects_invalid_edit_operations(studio, tmp_path):
    source = tmp_path / "contrat.docx"
    document = Document()
    document.add_paragraph("Texte")
    document.save(source)
    record, _ = studio.importer.import_file(source, source_kind="manual")

    result = await preview_document_edit_handler(None, record.id, "pas du json")
    assert not result.success
    assert "liste JSON valide" in result.output
