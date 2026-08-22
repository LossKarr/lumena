from __future__ import annotations

from pathlib import Path

from docx import Document
import openpyxl

from src.documents.document_edit_service import DocumentEditService
from src.documents.document_library import DocumentLibrary
from src.documents.import_service import DocumentImportService
from src.tools.document_hub import DocumentHub


def make_service(tmp_path):
    library = DocumentLibrary(tmp_path / "documents.sqlite")
    importer = DocumentImportService(tmp_path / "library", library)
    hub = DocumentHub(tmp_path / "outputs")
    return DocumentEditService(tmp_path / "versions", library=library, importer=importer, document_hub=hub), importer, library


def test_preview_does_not_mutate_original(tmp_path):
    service, importer, _ = make_service(tmp_path)
    source = tmp_path / "source.docx"
    doc = Document(); doc.add_paragraph("Bonjour ancien"); doc.save(source)
    record, _ = importer.import_file(source, source_kind="manual")
    before = source.read_bytes()
    preview = service.preview(record.id, [{"op": "replace_text", "find": "ancien", "replace": "nouveau"}])
    assert preview.format == "docx"
    assert source.read_bytes() == before


def test_docx_edit_creates_child_and_preserves_original(tmp_path):
    service, importer, library = make_service(tmp_path)
    source = tmp_path / "source.docx"
    doc = Document(); doc.add_paragraph("Bonjour ancien"); doc.save(source)
    record, _ = importer.import_file(source, source_kind="manual")
    result = service.apply(record.id, [{"op": "replace_text", "find": "ancien", "replace": "nouveau"}])
    child = library.get(result["record"]["id"])
    assert child.parent_id == record.id
    assert "nouveau" in " ".join(p.text for p in Document(child.path).paragraphs)
    assert "ancien" in " ".join(p.text for p in Document(source).paragraphs)
    assert result["operations_applied"] == 1


def test_xlsx_round_trip_preserves_unmodified_cell(tmp_path):
    service, importer, library = make_service(tmp_path)
    source = tmp_path / "book.xlsx"
    wb = openpyxl.Workbook(); ws = wb.active; ws["A1"] = "keep"; ws["B1"] = "old"; wb.save(source)
    record, _ = importer.import_file(source, source_kind="manual")
    result = service.apply(record.id, [{"op": "set_cell", "cell": "B1", "value": "new"}])
    out = openpyxl.load_workbook(library.get(result["record"]["id"]).path)
    assert out.active["A1"].value == "keep"
    assert out.active["B1"].value == "new"
