from __future__ import annotations

from pathlib import Path
import zipfile

import pytest

from src.documents.document_library import DocumentLibrary
from src.documents.document_security import DocumentSecurityError, inspect_document
from src.documents.import_service import DocumentImportService


@pytest.fixture
def service(tmp_path):
    library = DocumentLibrary(tmp_path / "studio" / "documents.sqlite")
    return DocumentImportService(tmp_path / "studio" / "library", library), library


def test_import_text_indexes_content_and_deduplicates(service, tmp_path):
    importer, library = service
    source = tmp_path / "rapport.txt"
    source.write_text("Analyse documentaire Lumena unique", encoding="utf-8")
    first, duplicate_1 = importer.import_file(source, source_kind="upload", source_uri="chat:web")
    second, duplicate_2 = importer.import_file(source, source_kind="upload", source_uri="chat:web")
    assert duplicate_1 is False
    assert duplicate_2 is True
    assert first.id == second.id
    assert Path(first.path).is_file()
    assert library.search("documentaire Lumena")[0].id == first.id


def test_import_docx_extracts_and_finds_content(service, tmp_path):
    from docx import Document

    source = tmp_path / "contrat.docx"
    doc = Document()
    doc.add_heading("Contrat Atlas", level=1)
    doc.add_paragraph("La clause albatros protège les livrables confidentiels de la mission.")
    doc.save(source)
    importer, library = service
    record, duplicate = importer.import_file(source, source_kind="manual")
    assert duplicate is False
    assert record.format == "docx"
    assert "albatros" in record.content_text
    assert library.search("albatros")[0].id == record.id


def test_extension_content_mismatch_is_rejected(tmp_path):
    fake = tmp_path / "fake.pdf"
    fake.write_text("not a pdf", encoding="utf-8")
    with pytest.raises(DocumentSecurityError, match="mismatch"):
        inspect_document(fake)


def test_macro_payload_in_ooxml_is_rejected(tmp_path):
    source = tmp_path / "bad.docx"
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("word/document.xml", "<doc/>")
        archive.writestr("word/vbaProject.bin", b"macro")
    with pytest.raises(DocumentSecurityError, match="macros"):
        inspect_document(source)


def test_archive_traversal_is_rejected(tmp_path):
    source = tmp_path / "bad.xlsx"
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("xl/workbook.xml", "<book/>")
        archive.writestr("../escape.txt", "bad")
    with pytest.raises(DocumentSecurityError, match="traversal"):
        inspect_document(source)


def test_library_format_filter(service, tmp_path):
    importer, library = service
    source = tmp_path / "notes.md"
    source.write_text("# Notes\n\nContenu filtrable pour test.", encoding="utf-8")
    importer.import_file(source, source_kind="manual")
    assert len(library.search("filtrable", formats=["md"])) == 1
    assert library.search("filtrable", formats=["pdf"]) == []


def test_fts_search_and_reference_resolution_accept_hyphenated_filename(service, tmp_path):
    importer, library = service
    source = tmp_path / "cert-doc-01.txt"
    source.write_text("Certificat documentaire Orion", encoding="utf-8")
    record, _ = importer.import_file(source, source_kind="manual")

    matches = library.search("cert-doc-01")
    assert [item.id for item in matches] == [record.id]
    assert library.resolve_reference("cert-doc-01.txt").id == record.id
    assert library.resolve_reference(record.id).id == record.id
